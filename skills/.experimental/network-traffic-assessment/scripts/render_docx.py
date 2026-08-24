#!/usr/bin/env python3
"""Genera el informe de Health Check / Assessment de firewall Palo Alto en .docx a partir de
un JSON estructurado, aplicando la paleta visual extraida del PDF de referencia de Palo Alto
Networks (ver reference/diagram-palette.md).

Uso:
    python3 render_docx.py informe.json Informe_Cliente_Dispositivo_SEK.docx

Formato del JSON de entrada: ver reference/render-docx-schema.md.
Requiere: pip install python-docx  (Pillow ya suele venir instalado; si no, pip install Pillow)
"""
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont

# ---------- Paleta Palo Alto Networks (ver reference/diagram-palette.md) ----------
# Valores extraidos directamente del contenido vectorial del PDF de referencia (PyMuPDF), no
# aproximados a ojo.
GOLD = RGBColor(0xFF, 0xCB, 0x06)          # nivel 1: titulos principales, portada
GOLD_HEX = "FFCB06"
ORANGE = RGBColor(0xFA, 0x58, 0x2D)        # nivel 2: categorias, severidad Critico/Alto
BLUE = RGBColor(0x4F, 0x81, 0xBD)          # nivel 3: nombre de cada check individual
GRAY_LABEL = RGBColor(0x80, 0x7C, 0x7B)    # etiquetas de campo (Findings/Recommendations/...)
LINK_BLUE = RGBColor(0x11, 0x55, 0xCC)     # enlaces de Referencias
ZEBRA_GRAY_HEX = "F2F2F2"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEXT_DARK = RGBColor(0x33, 0x33, 0x33)

HEADING_FONT = "Montserrat"
BODY_FONT = "Tahoma"

# Distincion visual por severidad en la lista de Recomendaciones priorizada (unica seccion del
# PDF de origen donde vale la pena diferenciar color por severidad - el resto del documento usa
# el mismo naranja para toda etiqueta de categoria). Solo colores ya verificados de la paleta.
SEVERITY_COLOR = {
    "CRÍTICO": ORANGE, "CRITICO": ORANGE,
    "ALTO": ORANGE,
    "IMPORTANTE": BLUE,
    "BAJO": GRAY_LABEL,
    "OTRAS RECOMENDACIONES": GRAY_LABEL,
}
SEVERITY_ORDER = ["Crítico", "Alto", "Importante", "Bajo", "Otras Recomendaciones"]

DIAGRAM_COLORS = {
    "internet_fill": (222, 234, 246), "internet_stroke": (0x4F, 0x81, 0xBD),
    "wan_fill": (255, 235, 200), "wan_stroke": (0xFA, 0x58, 0x2D),
    "fw_fill": (20, 20, 20),
    "lan_fill": (222, 234, 246), "lan_stroke": (0x4F, 0x81, 0xBD),
    "core_fill": (255, 224, 214), "core_stroke": (0xFA, 0x58, 0x2D),
    "vpn_fill": (235, 235, 234), "vpn_stroke": (0x80, 0x7C, 0x7B),
}


def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _style_table(table, zebra=True):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = BODY_FONT
        if i == 0:
            for cell in row.cells:
                _set_cell_bg(cell, GOLD_HEX)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = WHITE
        elif zebra and i % 2 == 0:
            for cell in row.cells:
                _set_cell_bg(cell, ZEBRA_GRAY_HEX)


def add_table(doc, headers, rows, severity_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if val is None else str(val)
    _style_table(table)

    if severity_col is not None:
        for ri, row in enumerate(rows, start=1):
            sev = str(row[severity_col]).strip().upper()
            color = SEVERITY_COLOR.get(sev)
            if color:
                for p in table.rows[ri].cells[severity_col].paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = color
    return table


def render_diagram_png(diagram, out_path, width=1400, height=800):
    """Redibuja el diagrama de arquitectura (Internet -> WAN -> Firewall -> LAN -> nucleo)
    con Pillow, usando la paleta de Palo Alto Networks (ver diagram-palette.md)."""
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_bold = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 16)
        font_reg = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 12)
    except Exception:
        font_bold = ImageFont.load_default()
        font_reg = ImageFont.load_default()

    cx = width // 2

    def box(x0, y0, x1, y1, fill, outline, label, sub=None, label_color=(0, 0, 0)):
        draw.rounded_rectangle([x0, y0, x1, y1], radius=8, fill=fill, outline=outline, width=2)
        tw = draw.textlength(label, font=font_bold)
        draw.text((x0 + (x1 - x0 - tw) / 2, y0 + 10), label, fill=label_color, font=font_bold)
        if sub:
            tw2 = draw.textlength(sub, font=font_reg)
            draw.text((x0 + (x1 - x0 - tw2) / 2, y0 + 32), sub, fill=(80, 80, 80), font=font_reg)

    # Internet
    draw.ellipse([cx - 160, 20, cx + 160, 100], fill=DIAGRAM_COLORS["internet_fill"],
                 outline=DIAGRAM_COLORS["internet_stroke"], width=2)
    draw.text((cx - 40, 50), "INTERNET", fill=DIAGRAM_COLORS["internet_stroke"], font=font_bold)

    # Panel de peers VPN (izquierda)
    peers = diagram.get("vpn_peers", [])
    if peers:
        ph = 40 + 18 * min(len(peers), 8)
        draw.rounded_rectangle([20, 130, 320, 130 + ph], radius=6,
                                fill=DIAGRAM_COLORS["vpn_fill"], outline=DIAGRAM_COLORS["vpn_stroke"], width=2)
        draw.text((40, 140), "Peers VPN IPSec", fill=DIAGRAM_COLORS["vpn_stroke"], font=font_bold)
        for i, p in enumerate(peers[:8]):
            draw.text((40, 165 + i * 18), f"- {p.get('name', '')} - {p.get('peer', '')}",
                       fill=(50, 50, 50), font=font_reg)

    # Fila WAN
    wan = diagram.get("wan", [])
    n = max(len(wan), 1)
    bw = min(260, max(120, (width - 80) // n - 20))
    start_x = cx - (n * bw + (n - 1) * 20) // 2
    wan_y = 170
    for i, w in enumerate(wan):
        x0 = start_x + i * (bw + 20)
        draw.line([cx, 100, x0 + bw // 2, wan_y], fill=DIAGRAM_COLORS["wan_stroke"], width=2)
        box(x0, wan_y, x0 + bw, wan_y + 60, DIAGRAM_COLORS["wan_fill"], DIAGRAM_COLORS["wan_stroke"],
            w.get("name", ""), " . ".join(w.get("ips", [])), DIAGRAM_COLORS["wan_stroke"])
        draw.line([x0 + bw // 2, wan_y + 60, x0 + bw // 2, 320], fill=DIAGRAM_COLORS["wan_stroke"], width=1)

    # Firewall
    fw_y0, fw_y1 = 320, 410
    draw.rounded_rectangle([cx - 190, fw_y0, cx + 190, fw_y1], radius=10, fill=DIAGRAM_COLORS["fw_fill"])
    fw_label = f"FIREWALL {diagram.get('model', '')}"
    draw.text((cx - draw.textlength(fw_label, font=font_bold) / 2, fw_y0 + 25), fw_label,
               fill="white", font=font_bold)
    fw_sub = f"{diagram.get('device', '')} . {diagram.get('os_label', 'PAN-OS')} {diagram.get('os_version', '')}"
    draw.text((cx - draw.textlength(fw_sub, font=font_reg) / 2, fw_y0 + 55), fw_sub,
               fill=(200, 200, 200), font=font_reg)

    # Fila LAN (+ nucleo con bypass si aplica)
    lan = diagram.get("lan", [])
    n2 = max(len(lan), 1)
    bw2 = min(260, max(120, (width - 80) // n2 - 20))
    start_x2 = cx - (n2 * bw2 + (n2 - 1) * 20) // 2
    lan_y = 500
    core = diagram.get("core")
    for i, l in enumerate(lan):
        x0 = start_x2 + i * (bw2 + 20)
        draw.line([cx, fw_y1, x0 + bw2 // 2, lan_y], fill=DIAGRAM_COLORS["lan_stroke"], width=2)
        box(x0, lan_y, x0 + bw2, lan_y + 60, DIAGRAM_COLORS["lan_fill"], DIAGRAM_COLORS["lan_stroke"],
            l.get("name", ""), " . ".join(l.get("ips", [])), DIAGRAM_COLORS["lan_stroke"])
        if core and l.get("is_core_transit"):
            cxm = x0 + bw2 // 2
            draw.line([cxm, lan_y + 60, cxm, lan_y + 80], fill=DIAGRAM_COLORS["core_stroke"], width=2)
            core_w = max(bw2, 220)
            cx0 = cxm - core_w // 2
            box(cx0, lan_y + 80, cx0 + core_w, lan_y + 170, DIAGRAM_COLORS["core_fill"],
                DIAGRAM_COLORS["core_stroke"], f"Nucleo {core.get('next_hop', '')}",
                f"{core.get('segment_count', '?')} subredes internas", DIAGRAM_COLORS["core_stroke"])
            draw.text((cx0 + 10, lan_y + 130), "Este-Oeste sin inspeccion (bypass)",
                       fill=DIAGRAM_COLORS["core_stroke"], font=font_reg)

    img.save(out_path)
    return out_path


def add_hyperlink(paragraph, url, text):
    """Inserta un hipervinculo real (no solo texto azul) en un parrafo de python-docx."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run_el.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run_el.append(text_el)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)
    return hyperlink


def _label_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = GRAY_LABEL
    return run


def _add_checks(doc, checks):
    """Renderiza el patron Device/Observacion -> Findings -> Recommendations -> References
    (ver SKILL.md). Cada check con hallazgo relevante es su propio bloque."""
    for c in checks or []:
        h3 = doc.add_heading(c.get("title", ""), level=3)
        for r in h3.runs:
            r.font.color.rgb = BLUE
            r.font.size = Pt(12)
            r.font.name = HEADING_FONT

        if c.get("device_observation"):
            p = doc.add_paragraph()
            _label_run(p, "Device / Observación: ")
            p.add_run(c["device_observation"])

        dt = c.get("data_table")
        if dt and dt.get("rows"):
            add_table(doc, dt["headers"], dt["rows"], severity_col=dt.get("severity_col"))
            doc.add_paragraph()

        if c.get("findings"):
            p = doc.add_paragraph()
            _label_run(p, "Findings:")
            for f in c["findings"]:
                doc.add_paragraph(f, style="List Bullet")

        if c.get("recommendations"):
            p = doc.add_paragraph()
            _label_run(p, "Recommendations:")
            for rec in c["recommendations"]:
                doc.add_paragraph(rec, style="List Bullet")

        refs = c.get("references")
        if refs:
            p = doc.add_paragraph()
            _label_run(p, "Referencias: ")
            for i, ref in enumerate(refs):
                if i > 0:
                    p.add_run("  |  ")
                url = ref.get("url", "")
                label = ref.get("label", url)
                if url:
                    add_hyperlink(p, url, label)
                else:
                    p.add_run(label)

        doc.add_paragraph()


def _add_recommendations_by_severity(doc, block):
    """Bloque de Recomendaciones priorizadas: bullets agrupados por severidad, sin tabla ni
    numeracion H-N (patron del PDF de referencia). `block` es un dict severidad -> lista de
    items; cada item es un string o {"text":..., "sub_bullets":[...]}."""
    for sev in SEVERITY_ORDER:
        items = block.get(sev)
        if not items:
            continue
        p = doc.add_paragraph()
        run = p.add_run(sev)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = SEVERITY_COLOR.get(sev.upper(), TEXT_DARK)
        for item in items:
            if isinstance(item, dict):
                doc.add_paragraph(item.get("text", ""), style="List Bullet")
                for sb in item.get("sub_bullets", []):
                    doc.add_paragraph(sb, style="List Bullet 2")
            else:
                doc.add_paragraph(str(item), style="List Bullet")
        doc.add_paragraph()


def _add_narrative(doc, narrative):
    if not narrative:
        return
    paras = narrative if isinstance(narrative, list) else [narrative]
    for para in paras:
        doc.add_paragraph(para)


def _add_tables(doc, tables):
    for t in tables or []:
        if t.get("caption"):
            p = doc.add_paragraph()
            r = p.add_run(t["caption"])
            r.font.italic = True
        add_table(doc, t["headers"], t["rows"], severity_col=t.get("severity_col"))
        doc.add_paragraph()


def _set_default_fonts(doc):
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10)


def _add_heading(doc, text, level, color):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = color
        r.font.name = HEADING_FONT
    return h


def _add_toc_field(doc):
    """Indice nativo de Word (campo TOC ligado a los estilos Heading 1-3), igual que el
    "Contents" del PDF de referencia. Word lo puebla al abrir el documento (o con F9/clic
    derecho > Actualizar campo) a partir de los add_heading() ya emitidos - no se arma a mano."""
    _add_heading(doc, "Índice", 1, GOLD)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    r_element = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Haga clic derecho y elija «Actualizar campo» para generar el índice."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    for el in (fld_begin, instr, fld_separate, placeholder, fld_end):
        r_element.append(el)

    # Pide a Word que actualice los campos (el TOC) al abrir el documento.
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    doc.add_page_break()


def build_docx(data, out_path):
    doc = Document()
    _set_default_fonts(doc)

    meta = data.get("meta", {})

    # Portada: Prepared for / Date / Prepared by / Version number (patron del PDF de referencia)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta.get("report_title", "Health Check / Assessment de Firewall"))
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = GOLD
    run.font.name = HEADING_FONT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for label, value in (
        ("Prepared for: ", meta.get("client", "")),
        ("Date: ", meta.get("date", "")),
        ("Prepared by: ", meta.get("prepared_by", "")),
        ("Version number: ", meta.get("version", "1.0")),
    ):
        p = doc.add_paragraph()
        _label_run(p, label)
        p.add_run(str(value))

    if meta.get("disclaimer"):
        doc.add_paragraph()
        h = _add_heading(doc, "Notices / Disclaimer", 2, GOLD)
        doc.add_paragraph(meta["disclaimer"])

    doc.add_page_break()
    _add_toc_field(doc)

    # Scope (tabla de dispositivos revisados)
    scope = data.get("scope")
    if scope:
        _add_heading(doc, "Scope", 1, GOLD)
        _add_narrative(doc, scope.get("narrative"))
        if scope.get("table"):
            add_table(doc, scope["table"]["headers"], scope["table"]["rows"])
        doc.add_paragraph()

    # Recommendations priorizadas por severidad
    recs = data.get("recommendations")
    if recs:
        _add_heading(doc, "Recommendations", 1, GOLD)
        _add_recommendations_by_severity(doc, recs)

    # Secciones / categorias principales
    for section in data.get("sections", []):
        number = section.get("number")
        title_txt = f"{number}. {section['title']}" if number else section["title"]
        _add_heading(doc, title_txt, 1, GOLD)

        _add_narrative(doc, section.get("narrative"))

        if section.get("diagram"):
            base = Path(out_path).with_suffix("")
            png_path = f"{base}_diagrama.png"
            render_diagram_png(section["diagram"], png_path)
            doc.add_picture(png_path, width=Inches(6.3))

        _add_tables(doc, section.get("tables"))
        _add_checks(doc, section.get("checks"))

        for sub_s in section.get("subsections", []):
            _add_heading(doc, sub_s.get("title", ""), 2, ORANGE)
            _add_narrative(doc, sub_s.get("narrative"))
            if sub_s.get("diagram"):
                base = Path(out_path).with_suffix("")
                png_path = f"{base}_diagrama.png"
                render_diagram_png(sub_s["diagram"], png_path)
                doc.add_picture(png_path, width=Inches(6.3))
            _add_tables(doc, sub_s.get("tables"))
            _add_checks(doc, sub_s.get("checks"))

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) != 3:
        print("Uso: python3 render_docx.py <informe.json> <salida.docx>", file=sys.stderr)
        sys.exit(1)
    json_path, out_path = sys.argv[1], sys.argv[2]
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    build_docx(data, out_path)
    print(f"Generado: {out_path}")


if __name__ == "__main__":
    main()
